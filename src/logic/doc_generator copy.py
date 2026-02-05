from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import re

def create_word_docx(markdown_text):
    """แปลง Markdown Text ให้กลายเป็นไฟล์ Word สวยๆ"""
    doc = Document()
    
    # ตั้งค่า Font (ถ้าเครื่องมี TH SarabunNew จะดีมาก แต่ใช้ Default ไปก่อน)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Tahoma'
    font.size = Pt(11)

    # วนลูปอ่านทีละบรรทัดเพื่อจัดรูปแบบ
    for line in markdown_text.split('\n'):
        line = line.strip()
        
        if not line: continue # ข้ามบรรทัดว่าง

        # จัดการหัวข้อ (H1, H2)
        if line.startswith('# '):
            p = doc.add_heading(line.replace('# ', ''), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=1)
        
        # จัดการตัวหนา (Bold) **...**
        elif line.startswith('**') or line.startswith('* '):
            p = doc.add_paragraph()
            clean_text = line.replace('**', '').replace('* ', '• ')
            runner = p.add_run(clean_text)
            # ถ้าเป็น Bullet ให้ทำตัวธรรมดา แต่ถ้าเป็น Header ย่อยให้ทำตัวหนา
            if not line.startswith('* '):
                runner.bold = True
        
        # จัดการตาราง (Table) แบบง่าย
        elif '|' in line and '---' not in line:
            # ถ้าเจอเส้นตาราง ให้เขียนเป็น text ธรรมดาไปก่อน (เพราะการสร้าง Table ใน Word ซับซ้อน)
            # แต่เราจะจัด format ให้ดูง่าย
            clean_line = line.replace('|', '  ').strip()
            p = doc.add_paragraph(clean_line)
            p.paragraph_format.left_indent = Pt(20) # ย่อหน้าเข้าไปหน่อย
        
        # ข้อความทั่วไป
        elif '---' not in line: # ข้ามเส้นขีดคั่น
            doc.add_paragraph(line)

    # บันทึกลง Memory Buffer (ไม่ได้เซฟลงเครื่องโดยตรง)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer