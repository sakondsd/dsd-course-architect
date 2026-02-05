import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, font_name='TH Sarabun New', size=16, bold=False):
    """ฟังก์ชันช่วยตั้งค่าฟอนต์ภาษาไทย"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def process_markdown_text(paragraph, text):
    """
    ฟังก์ชันช่วยแปลง Markdown (**ตัวหนา**) และ <br> ให้เป็น Format ของ Word
    """
    # 1. แปลง <br> เป็น Newline (\n)
    text = text.replace("<br>", "\n").replace("<br/>", "\n")
    
    # 2. แยกส่วนข้อความด้วยตัวหนา **...**
    # Regex นี้จะแยกข้อความออกเป็นกลุ่มๆ เช่น ['ปกติ', '**ตัวหนา**', 'ปกติ']
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # เป็นตัวหนา ตัด ** ออก
            clean_text = part[2:-2]
            run = paragraph.add_run(clean_text)
            set_font(run, bold=True)
        else:
            # เป็นตัวธรรมดา
            if part: # ไม่เอาว่างเปล่า
                run = paragraph.add_run(part)
                set_font(run, bold=False)

def create_word_docx(markdown_content):
    doc = Document()
    
    # --- 1. ตั้งค่าหน้ากระดาษ ---
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4 Width
    section.page_height = Inches(11.69) # A4 Height
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- 2. อ่านทีละบรรทัดเพื่อแปลงผล ---
    lines = markdown_content.split('\n')
    
    table_mode = False
    table_data = []
    
    for line in lines:
        line = line.strip()
        
        # --- A. ตรวจจับว่าเป็นตารางหรือไม่? ---
        if line.startswith('|') and line.endswith('|'):
            table_mode = True
            # ตัด | หัวท้ายออก และแบ่งช่องด้วย |
            cells = [c.strip() for c in line.split('|')[1:-1]]
            
            # ข้ามบรรทัดเส้นขีด (เช่น |---|---|)
            if set(line.replace('|', '').replace('-', '').replace(':', '').strip()) == set():
                continue
                
            table_data.append(cells)
            continue
        
        # --- B. ถ้าจบโซนตารางแล้ว ให้วาดตารางลง Word ---
        if table_mode and (not line.startswith('|')):
            table_mode = False
            if table_data:
                # สร้างตารางใน Word
                rows = len(table_data)
                cols = len(table_data[0])
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid' # ใส่เส้นตาราง
                
                for r_idx, row_data in enumerate(table_data):
                    row = table.rows[r_idx]
                    for c_idx, cell_text in enumerate(row_data):
                        # ป้องกันกรณีข้อมูลไม่ครบช่อง
                        if c_idx < len(row.cells):
                            cell = row.cells[c_idx]
                            # ล้าง paragraph เดิมที่ว่างเปล่าออก
                            cell._element.clear_content()
                            p = cell.add_paragraph()
                            
                            # ถ้าเป็นหัวตาราง (แถวแรก) ให้ทำตัวหนาและจัดกึ่งกลาง
                            if r_idx == 0:
                                run = p.add_run(cell_text)
                                set_font(run, bold=True)
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                # ใส่สีพื้นหลังหัวตาราง (สีเทาอ่อน)
                                tcPr = cell._element.get_or_add_tcPr()
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:val'), 'clear')
                                shd.set(qn('w:color'), 'auto')
                                shd.set(qn('w:fill'), 'E0E0E0') # Hex Color
                                tcPr.append(shd)
                            else:
                                # เนื้อหาในตาราง (รองรับตัวหนาและ <br>)
                                process_markdown_text(p, cell_text)
                                
                # เว้นบรรทัดหลังตารางนิดนึง
                doc.add_paragraph()
                table_data = [] # รีเซ็ตข้อมูลตาราง

        # --- C. จัดการข้อความปกติ (Heading / List / Text) ---
        if not table_mode:
            if line.startswith('# '):
                p = doc.add_heading(line.replace('# ', ''), level=1)
                for run in p.runs: set_font(run, size=24, bold=True)
            elif line.startswith('## '):
                p = doc.add_heading(line.replace('## ', ''), level=2)
                for run in p.runs: set_font(run, size=20, bold=True)
            elif line.startswith('* ') or line.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                process_markdown_text(p, line[2:])
            else:
                if line.strip(): # ไม่เอาบรรทัดว่างเกินไป
                    p = doc.add_paragraph()
                    process_markdown_text(p, line)

    # Save ลง Memory Buffer
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io